"""文档解析流水线。

解析层只负责把 HTML/JSON/PDF/DOCX/XLSX 变成带页级信息的文本和表格，
不负责决定项目状态。固定来源走规则解析；只有显式传入 ``use_crawl4ai``
时才启动 Crawl4AI。
"""

from __future__ import annotations

import asyncio
import json
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Iterable

from bs4 import BeautifulSoup

from tender_ai.documents.quality import document_quality_score, extract_pdf
from tender_ai.urls import content_hash


@dataclass(frozen=True)
class DocumentPage:
    page_number: int
    text: str
    tables: tuple[tuple[str, ...], ...] = ()


@dataclass(frozen=True)
class ParsedDocument:
    source_url: str | None
    source_file: str | None
    content_type: str
    parser: str
    text: str
    pages: tuple[DocumentPage, ...] = ()
    tables: tuple[tuple[str, ...], ...] = ()
    quality_score: float = 0.0
    used_ocr: bool = False
    used_mineru: bool = False
    error: str | None = None
    content_hash: str = ""
    metadata: dict[str, Any] = field(default_factory=dict)

    @property
    def page_count(self) -> int:
        return len(self.pages)

    def page_for(self, fragment: str | None) -> int | None:
        if not fragment:
            return None
        needle = re.sub(r"\s+", "", str(fragment))
        if not needle:
            return None
        for page in self.pages:
            if needle in re.sub(r"\s+", "", page.text):
                return page.page_number
        return None


def _clean_text(value: Any) -> str:
    text = "" if value is None else str(value)
    text = text.replace("\x00", "")
    text = re.sub(r"[ \t\f\v]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _decode(data: bytes) -> str:
    for encoding in ("utf-8-sig", "gb18030", "big5"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")


def _table_rows(soup: BeautifulSoup) -> tuple[tuple[str, ...], ...]:
    rows: list[tuple[str, ...]] = []
    for table in soup.select("table"):
        for row in table.select("tr"):
            cells = tuple(_clean_text(cell.get_text(" ", strip=True)) for cell in row.select("th,td"))
            if any(cells):
                rows.append(cells)
    return tuple(rows)


def _html_text(soup: BeautifulSoup, selector: str | None = None) -> str:
    for node in soup.select("script,style,noscript,template,svg,canvas"):
        node.decompose()
    candidates = []
    if selector:
        candidates.extend(soup.select(selector))
    if not candidates:
        for candidate_selector in (
            "article", "main", ".article", ".article-content", ".notice-content",
            ".detail-content", ".content", "#content", ".正文", ".内容",
        ):
            candidates.extend(soup.select(candidate_selector))
    node = max(candidates, key=lambda item: len(item.get_text(" ", strip=True))) if candidates else soup
    return _clean_text(node.get_text("\n", strip=True))


def _parsed(
    *,
    source_url: str | None,
    source_file: str | None,
    content_type: str,
    parser: str,
    text: str,
    pages: Iterable[DocumentPage] = (),
    tables: Iterable[tuple[str, ...]] = (),
    expected_terms: tuple[str, ...] = (),
    used_ocr: bool = False,
    used_mineru: bool = False,
    error: str | None = None,
    metadata: dict[str, Any] | None = None,
) -> ParsedDocument:
    clean = _clean_text(text)
    page_rows = tuple(pages)
    if not page_rows and clean:
        page_rows = (DocumentPage(1, clean),)
    return ParsedDocument(
        source_url=source_url,
        source_file=source_file,
        content_type=content_type,
        parser=parser,
        text=clean,
        pages=page_rows,
        tables=tuple(tables),
        quality_score=document_quality_score(clean, expected_terms=expected_terms),
        used_ocr=used_ocr,
        used_mineru=used_mineru,
        error=error,
        content_hash=content_hash(clean),
        metadata=metadata or {},
    )


def parse_html(
    html: str,
    *,
    source_url: str | None = None,
    source_file: str | None = None,
    selector: str | None = None,
    expected_terms: tuple[str, ...] = (),
    use_crawl4ai: bool = False,
) -> ParsedDocument:
    """规则优先清洗 HTML；复杂陌生页面只有显式允许时才调用 Crawl4AI。"""

    soup = BeautifulSoup(html or "", "lxml")
    text = _html_text(soup, selector)
    tables = _table_rows(soup)
    result = _parsed(
        source_url=source_url,
        source_file=source_file,
        content_type="text/html",
        parser="html.selector_rule" if selector else "html.beautifulsoup_rule",
        text=text,
        pages=(DocumentPage(1, text, tables),) if text else (),
        tables=tables,
        expected_terms=expected_terms,
    )
    if not use_crawl4ai or result.quality_score >= 35 or not source_url:
        return result
    try:
        crawl_text = _crawl4ai_fetch(source_url)
    except Exception as exc:  # pragma: no cover - 只有显式启用浏览器才进入
        return ParsedDocument(**{**result.__dict__, "error": f"crawl4ai:{exc}"})
    browser_result = _parsed(
        source_url=source_url,
        source_file=source_file,
        content_type="text/html",
        parser="crawl4ai",
        text=crawl_text,
        pages=(DocumentPage(1, crawl_text),) if crawl_text else (),
        expected_terms=expected_terms,
    )
    return browser_result if browser_result.quality_score > result.quality_score else result


def _crawl4ai_fetch(url: str) -> str:
    async def _run() -> str:
        from crawl4ai import AsyncWebCrawler

        async with AsyncWebCrawler() as crawler:
            result = await crawler.arun(url=url)
            markdown = getattr(result, "markdown", None)
            if isinstance(markdown, str):
                return markdown
            return str(getattr(result, "cleaned_html", "") or "")

    try:
        asyncio.get_running_loop()
    except RuntimeError:
        return asyncio.run(_run())
    raise RuntimeError("当前事件循环中不能同步启动 Crawl4AI")


def parse_json_text(
    text: str,
    *,
    source_url: str | None = None,
    source_file: str | None = None,
    expected_terms: tuple[str, ...] = (),
) -> ParsedDocument:
    try:
        value = json.loads(text or "")
    except json.JSONDecodeError:
        return _parsed(source_url=source_url, source_file=source_file, content_type="application/json", parser="json.fallback_text", text=text, expected_terms=expected_terms, error="JSON_PARSE_ERROR")

    lines: list[str] = []

    def walk(item: Any, prefix: str = "") -> None:
        if isinstance(item, dict):
            for key, value_item in item.items():
                walk(value_item, f"{prefix}.{key}" if prefix else str(key))
        elif isinstance(item, list):
            for index, value_item in enumerate(item):
                walk(value_item, f"{prefix}[{index}]")
        else:
            lines.append(f"{prefix}: {item}")

    walk(value)
    return _parsed(source_url=source_url, source_file=source_file, content_type="application/json", parser="json.rule", text="\n".join(lines), expected_terms=expected_terms)


def parse_docx(path: str | Path, *, source_url: str | None = None, expected_terms: tuple[str, ...] = ()) -> ParsedDocument:
    target = Path(path)
    try:
        from docx import Document

        document = Document(str(target))
        paragraphs = [_clean_text(item.text) for item in document.paragraphs if _clean_text(item.text)]
        tables: list[tuple[str, ...]] = []
        for table in document.tables:
            for row in table.rows:
                cells = tuple(_clean_text(cell.text) for cell in row.cells)
                if any(cells):
                    tables.append(cells)
        text = "\n".join(paragraphs + [" | ".join(row) for row in tables])
        return _parsed(source_url=source_url, source_file=str(target), content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", parser="python-docx", text=text, pages=(DocumentPage(1, text, tuple(tables)),) if text else (), tables=tables, expected_terms=expected_terms)
    except Exception as exc:
        return _parsed(source_url=source_url, source_file=str(target), content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", parser="python-docx.error", text="", expected_terms=expected_terms, error=str(exc))


def parse_xlsx(path: str | Path, *, source_url: str | None = None, expected_terms: tuple[str, ...] = ()) -> ParsedDocument:
    target = Path(path)
    try:
        from openpyxl import load_workbook

        workbook = load_workbook(str(target), read_only=True, data_only=True)
        page_rows: list[DocumentPage] = []
        all_tables: list[tuple[str, ...]] = []
        for page_number, sheet in enumerate(workbook.worksheets, start=1):
            rows: list[str] = []
            for row in sheet.iter_rows(values_only=True):
                cells = tuple(_clean_text(value) for value in row)
                if any(cells):
                    all_tables.append(cells)
                    rows.append(" | ".join(cells))
            page_rows.append(DocumentPage(page_number, "\n".join(rows), tuple(all_tables[-len(rows):]) if rows else ()))
        text = "\n\n".join(f"[{sheet.title}]\n{page.text}" for sheet, page in zip(workbook.worksheets, page_rows) if page.text)
        workbook.close()
        return _parsed(source_url=source_url, source_file=str(target), content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", parser="openpyxl", text=text, pages=page_rows, tables=all_tables, expected_terms=expected_terms)
    except Exception as exc:
        return _parsed(source_url=source_url, source_file=str(target), content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", parser="openpyxl.error", text="", expected_terms=expected_terms, error=str(exc))


def parse_pdf_document(path: str | Path, *, source_url: str | None = None, expected_terms: tuple[str, ...] = ()) -> ParsedDocument:
    target = Path(path)
    quality = extract_pdf(target, expected_terms=expected_terms)
    mineru_document: ParsedDocument | None = None
    if quality.needs_mineru:
        mineru_document = parse_mineru(target, source_url=source_url, expected_terms=expected_terms)
        if mineru_document.text and mineru_document.quality_score > quality.score:
            return ParsedDocument(**{**mineru_document.__dict__, "used_mineru": True, "metadata": {**mineru_document.metadata, "pymupdf4llm_quality_score": quality.score, "mineru_fallback": True}})
    pages = tuple(DocumentPage(index, page_text) for index, page_text in enumerate(quality.pages, start=1) if page_text)
    return ParsedDocument(
        source_url=source_url,
        source_file=str(target),
        content_type="application/pdf",
        parser=quality.parser,
        text=_clean_text(quality.text),
        pages=pages,
        quality_score=quality.score,
        used_ocr=quality.used_ocr,
        error=quality.error,
        content_hash=content_hash(quality.text),
        metadata={"needs_mineru": quality.needs_mineru, "mineru_fallback_attempted": mineru_document is not None, "mineru_error": mineru_document.error if mineru_document is not None else None},
    )


def parse_mineru(path: str | Path, *, source_url: str | None = None, expected_terms: tuple[str, ...] = ()) -> ParsedDocument:
    """MinerU fallback边界。

    MinerU 是可选重型依赖。安装后可在这里接入其稳定 CLI/API；未安装时明确
    返回失败状态，调用方仍保留 PyMuPDF4LLM 的文本和证据，不会伪造 OCR 结果。
    """

    try:
        import mineru  # type: ignore  # pragma: no cover

        _ = mineru
    except ImportError:
        return _parsed(source_url=source_url, source_file=str(path), content_type="application/pdf", parser="mineru.unavailable", text="", expected_terms=expected_terms, error="MINERU_NOT_INSTALLED")
    return _parsed(source_url=source_url, source_file=str(path), content_type="application/pdf", parser="mineru.adapter_pending", text="", expected_terms=expected_terms, error="MINERU_API_NOT_CONFIGURED")


def parse_path(path: str | Path, *, source_url: str | None = None, content_type: str | None = None, expected_terms: tuple[str, ...] = ()) -> ParsedDocument:
    target = Path(path)
    suffix = target.suffix.lower()
    detected = (content_type or "").lower()
    if suffix in {".html", ".htm"} or "html" in detected:
        return parse_html(_decode(target.read_bytes()), source_url=source_url, source_file=str(target), expected_terms=expected_terms)
    if suffix == ".json" or "json" in detected:
        return parse_json_text(_decode(target.read_bytes()), source_url=source_url, source_file=str(target), expected_terms=expected_terms)
    if suffix == ".pdf" or "pdf" in detected:
        return parse_pdf_document(target, source_url=source_url, expected_terms=expected_terms)
    if suffix == ".docx" or "wordprocessingml" in detected:
        return parse_docx(target, source_url=source_url, expected_terms=expected_terms)
    if suffix in {".xlsx", ".xlsm"} or "spreadsheetml" in detected or "excel" in detected:
        return parse_xlsx(target, source_url=source_url, expected_terms=expected_terms)
    return _parsed(source_url=source_url, source_file=str(target), content_type=detected or "application/octet-stream", parser="unsupported", text="", expected_terms=expected_terms, error="UNSUPPORTED_DOCUMENT_TYPE")


__all__ = ["DocumentPage", "ParsedDocument", "parse_docx", "parse_html", "parse_json_text", "parse_mineru", "parse_path", "parse_pdf_document", "parse_xlsx"]
